from docx import Document
import os
from datetime import datetime

def heading_solver(tag):
	header={"report_start": "Inicio da missão",
			"report_end": "Fim da missão",
			"GPS_LOW": "Alerta de GPS"}
	return header[tag]

def report_generator(Report):
	directory = 'reports'
	file_name = 'example.docx'
	file_path = os.path.join(directory, file_name)
	os.makedirs(directory, exist_ok=True)
	doc = Document()

	current_date = datetime.now().date()
	formatted_date = current_date.strftime('%Y/%m/%d')

	doc.add_heading(f'Relatório de Vôo {formatted_date}', 0)

	doc.add_heading('Objetivo', level=1)
	doc.add_paragraph("[Descrição da missão]")

	doc.add_heading('Missão')
	for point in Report:
		doc.add_heading(heading_solver(point["tag"]), level=2)

	# Add a Section with a Heading
	doc.add_heading('Section 1', level=1)
	doc.add_paragraph('Content for the first section.')

	# Add a Subsection with a Heading
	doc.add_heading('Subsection 1.1', level=2)
	doc.add_paragraph('Content for the first subsection.')

	# Save the document in the "reports" folder
	doc.save(file_path)

	print(f"Document saved in: {file_path}")